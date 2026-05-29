from docx import Document
from docx.shared import Inches
import os

doc = Document()
doc.add_heading("Assignment #4 — ThingsBoard Hello World", 0)
doc.add_paragraph("Student: Yakira_S")
doc.add_paragraph("Course: IoT")
doc.add_paragraph("Date: 2026-05-29")
doc.add_paragraph()

steps = [
    ("Login to ThingsBoard Cloud",
     "Signed up at thingsboard.cloud and logged in via GitHub OAuth. "
     "The ThingsBoard Tenant interface homepage confirms successful authentication.",
     "../screenshots/01-login.png"),
    ("Step 1: Device Provisioned",
     "Created a device named 'Thermometer' using Entities → Devices → Add new device. "
     "The device was created with Device profile: default.",
     "../screenshots/02-device-provisioned.png"),
    ("Step 1: Telemetry Active",
     "Used the built-in 'Check connectivity' tool to send temperature=25 via HTTP POST. "
     "Device status changed from Inactive to Active and Latest Telemetry shows temperature=25.",
     "../screenshots/03-telemetry-active.png"),
    ("Step 2: Real-Time Dashboard",
     "Created 'My Dashboard' with three widgets: Value card (current temperature), "
     "Time series chart (historical data), and Alarms table — all showing live data from the Thermometer device.",
     "../screenshots/04-dashboard.png"),
    ("Step 3: Alarm Rule Defined",
     "Created a 'High temperature' alarm rule via Alarms → Alarm rules. "
     "Condition: return temperature > 25; (Script mode, severity: Critical, Device profile: default).",
     "../screenshots/05-alarm-rule.png"),
    ("Step 3: Alarm Triggered",
     "Sent temperature=26 via Check connectivity (above the threshold of 25). "
     "ThingsBoard raised a Critical 'High temperature' alarm visible in the Alarm list.",
     "../screenshots/06-alarm-triggered.png"),
    ("Step 4: Customer Created",
     "Created customer 'My New Customer' and assigned the Thermometer device and My Dashboard to them "
     "via Manage owner and groups on each entity.",
     "../screenshots/07-customer-created.png"),
    ("Step 5: Customer User Activated",
     "Created customer user account (customer.test.iot@yopmail.com) and activated it via the "
     "Display activation link method. The customer user sees only Home, Resources, and Notification center — "
     "no admin features, confirming the limited customer role.",
     "../screenshots/08-customer-user-activated.png"),
]

for title, description, img_path in steps:
    doc.add_heading(title, level=1)
    doc.add_paragraph(description)
    abs_path = os.path.abspath(os.path.join(os.path.dirname(__file__), img_path))
    if os.path.exists(abs_path):
        doc.add_picture(abs_path, width=Inches(5.5))
    else:
        doc.add_paragraph(f"[Screenshot not found: {img_path}]")
    doc.add_paragraph()

out_path = os.path.join(os.path.dirname(__file__), "Assignment4_Report.docx")
doc.save(out_path)
print(f"Report saved to: {out_path}")
